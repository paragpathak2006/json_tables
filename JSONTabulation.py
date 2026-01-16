import json
from docx import Document
from Logger import Logger 
from html4docx import HtmlToDocx

"""
The main recursive tabulation class 
"""
class JSONTabulator:
    def __init__(self, doc: Document, data: dict | list):
        assert isinstance(data, (dict, list)), "Data must be a dictionary or a list."   
        Logger.log("Initialized JSONTabulator")
        self.doc = doc
        self.data = data
        self.html_converter = HtmlToDocx()

    """ The main tabulation function """
    def convert(self):
        Logger.log("Called JSONTabulator::convert()")
        self.tabulate(self.doc, self.data)

    """ The recursive tabulation function """
    def tabulate(self, container: Document, data: dict | list | str ):
        Logger.log("Called JSONTabulator::tabulate()")

        # Base case: if data is a string, add as paragraph
        if not isinstance(data, (dict, list)):
            self.html_converter.add_html_to_cell(str(data), container)
            return

        # if data is an array of objects, create excel-style table
        if self.is_arr(data):
            self.excel_table(container, data)
            return

        # if data is a dictionary , create a two-column table
        if isinstance(data, dict):
            table = container.add_table(rows=len(data), cols=2)
            table.style = "Table Grid"

            for row, (key, value) in enumerate(data.items()):
                key_cell = table.cell(row, 0)
                self.html_converter.add_html_to_cell(str(key), key_cell)

                value_cell = table.cell(row, 1)
                value_cell.text = ""

                if isinstance(value, (dict, list)):
                    self.tabulate(value_cell, value)
                else:
                    self.html_converter.add_html_to_cell("" if value is None else str(value), value_cell)

            return

        # if data is a list of values, create a single-column table
        if isinstance(data, list):
            table = container.add_table(rows=len(data), cols=1)
            table.style = "Table Grid"

            for i, item in enumerate(data):
                cell = table.cell(i, 0)
                cell.text = ""
                if isinstance(item, (dict, list)):
                    self.tabulate(cell, item)
                else:
                    self.html_converter.add_html_to_cell(str(item), cell)
                    
        return

    """ Excel-style table for array of objects """
    def excel_table(self, container: Document, array : list):
        Logger.log("Creating Excel-style table")
        Logger.log("Called JSONTabulator::excel_table()")

        # All columns 
        columns = []
        for obj in array:
            for key in obj.keys():
                if key not in columns:
                    columns.append(key)

        table = container.add_table(rows=len(array) + 1, cols=len(columns))
        table.style = "Table Grid"

        # First row
        for index, column in enumerate(columns):
            cell = table.cell(0, index)
            self.html_converter.add_html_to_cell(column, cell)

        # Data rows
        for row, obj in enumerate(array, start=1):
            for index, column in enumerate(columns):
                cell = table.cell(row, index)
                cell.text = ""

                value = obj.get(column)
                if isinstance(value, (dict, list)):
                    self.tabulate(cell, value)
                else:
                    self.html_converter.add_html_to_cell("" if value is None else str(value), cell)

    """checks if input is array of objects"""
    def is_arr(self, arr):
        return (
            isinstance(arr, list)
            and arr
            and all(isinstance(x, dict) for x in arr)
        )

    """
    The main recursive tabulation class 
    """

