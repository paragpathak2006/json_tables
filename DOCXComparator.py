from docx import Document
import zipfile
import xml.etree.ElementTree as ET

class DOCXComparator:
    
    @classmethod
    def compare_text_and_tables(cls, docx_path_1: str, docx_path_2: str) -> bool:
        content_1 = cls.extract_text_and_tables(docx_path_1)
        content_2 = cls.extract_text_and_tables(docx_path_2)

        return content_1 == content_2

    @classmethod
    def compare_document_xml(cls, docx_path_1: str, docx_path_2: str) -> bool:
        content_1 = cls.extract_document_xml(docx_path_1)
        content_2 = cls.extract_document_xml(docx_path_2)

        return content_1 == content_2

    @classmethod
    def extract_text_and_tables(cls, docx_path)-> list:
        doc = Document(docx_path)
        result = []

        # paragraphs
        for p in doc.paragraphs:
            result.append(("P", p.text.strip()))

        # tables
        for table in doc.tables:
            for row in table.rows:
                result.append(
                    ("T", [cell.text.strip() for cell in row.cells])
                )

        return result

    @classmethod
    def extract_document_xml(cls, docx_path) -> bytes:
        with zipfile.ZipFile(docx_path) as z:
            xml = z.read("word/document.xml")
        return ET.tostring(ET.fromstring(xml))


